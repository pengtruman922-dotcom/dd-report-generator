"""Extract text from Word (.docx) files.

Extracts text from paragraphs, tables, and embedded images (via OCR).
"""

from __future__ import annotations

import logging
from pathlib import Path

from parsers.ocr_utils import ocr_image

log = logging.getLogger(__name__)


def extract_docx_text(file_path) -> str:
    """Return the full text of a .docx file, including OCR of embedded images.

    Args:
        file_path: A file path (str/Path) or a file-like object (e.g. BytesIO).
    """
    try:
        from docx import Document
    except ImportError:
        log.warning("python-docx not installed, cannot parse .docx")
        return ""

    try:
        # python-docx Document() accepts both file paths and file-like objects
        if isinstance(file_path, (str, Path)):
            doc = Document(str(file_path))
        else:
            doc = Document(file_path)
    except Exception:
        log.exception("Failed to open .docx: %s", file_path)
        return ""

    parts: list[str] = []

    # 1. Extract text from paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # 2. Extract text from tables
    for table in doc.tables:
        rows_text: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            row_str = " | ".join(cells)
            if row_str.replace("|", "").strip():
                rows_text.append(row_str)
        if rows_text:
            parts.append("\n".join(rows_text))

    # 3. Extract and OCR embedded images
    image_count = 0
    ocr_parts: list[str] = []
    try:
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                img_bytes = rel.target_part.blob
                ocr_text = ocr_image(img_bytes)
                if ocr_text.strip():
                    ocr_parts.append(ocr_text)
                    image_count += 1
    except Exception:
        log.exception("Failed to extract images from .docx: %s", file_path)

    if ocr_parts:
        parts.append("\n\n".join(ocr_parts))
        log.info("DOCX: extracted text from %d paragraphs + %d images (OCR) in %s",
                 len(doc.paragraphs), image_count, file_path)

    return "\n\n".join(parts)
